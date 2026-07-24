"""
Resume Parser Module
Handles document ingestion and text extraction for PDF and DOCX formats.
"""

import io
import os
import pdfplumber
import docx
from src.config import config
from src.logger import logger
from src.utils import clean_text, is_valid_file_extension


class ParsingError(Exception):
    """Custom exception raised when file parsing fails."""
    pass


class UnsupportedFileTypeError(ParsingError):
    """Custom exception raised when an unsupported file type is provided."""
    pass


class ResumeParser:
    """Parser class responsible for extracting raw text from resume documents."""

    def parse_file(self, file_source: str | bytes, filename: str) -> str:
        """
        Parses a resume file and returns clean extracted text.

        Args:
            file_source: File path string or file bytes stream.
            filename: Original filename used to detect extension.

        Returns:
            str: Extracted and sanitized text content.

        Raises:
            UnsupportedFileTypeError: If file format is not supported (.pdf or .docx).
            ParsingError: If text extraction fails.
        """
        # If file_source is already raw text content (e.g. sample text or text input), return directly
        if isinstance(file_source, str) and not (file_source.endswith(".pdf") or file_source.endswith(".docx")) and len(file_source) > 50:
            cleaned = clean_text(file_source)
            if cleaned:
                return cleaned

        if not is_valid_file_extension(filename):
            logger.error(f"Unsupported file format attempted: {filename}")
            raise UnsupportedFileTypeError(
                f"Unsupported file format '{filename}'. Allowed formats: {config.ALLOWED_EXTENSIONS}"
            )

        ext = os.path.splitext(filename)[1].lower()
        logger.info(f"Parsing resume document: '{filename}' (type: {ext})")

        try:
            if ext == ".pdf":
                extracted_text = self._parse_pdf(file_source)
            elif ext == ".docx":
                extracted_text = self._parse_docx(file_source)
            else:
                raise UnsupportedFileTypeError(f"Unsupported extension: {ext}")

            cleaned = clean_text(extracted_text)
            if not cleaned:
                logger.warning(f"No text extracted from file: {filename}")
                raise ParsingError(
                    "Extracted text is empty. The file may be image-only or empty."
                )

            logger.info(
                f"Successfully extracted {len(cleaned)} characters from '{filename}'"
            )
            return cleaned

        except ParsingError:
            raise
        except Exception as e:
            logger.error(f"Failed to parse file '{filename}': {str(e)}", exc_info=True)
            raise ParsingError(f"Error parsing file '{filename}': {str(e)}") from e

    def _parse_pdf(self, file_source: str | bytes) -> str:
        """Extracts text content page-by-page from a PDF file using pdfplumber with PyPDF2 fallback."""
        extracted_pages: list[str] = []

        try:
            if isinstance(file_source, bytes):
                pdf_stream = io.BytesIO(file_source)
                with pdfplumber.open(pdf_stream) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text and text.strip():
                            extracted_pages.append(text)
            else:
                with pdfplumber.open(file_source) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text and text.strip():
                            extracted_pages.append(text)

            res = "\n".join(extracted_pages)
            if res.strip():
                return res

        except Exception as e:
            logger.warning(f"pdfplumber extraction encountered issue: {e}")

        # Fallback to PyPDF2
        try:
            import PyPDF2
            if isinstance(file_source, bytes):
                reader = PyPDF2.PdfReader(io.BytesIO(file_source))
            else:
                reader = PyPDF2.PdfReader(file_source)

            fallback_pages = []
            for page in reader.pages:
                t = page.extract_text()
                if t and t.strip():
                    fallback_pages.append(t)
            return "\n".join(fallback_pages)
        except Exception as fallback_err:
            logger.error(f"PyPDF2 fallback also failed: {fallback_err}")
            raise ParsingError(f"PDF parsing failure: {str(fallback_err)}") from fallback_err

    def _parse_docx(self, file_source: str | bytes) -> str:
        """Extracts text content from paragraphs and tables in DOCX files using python-docx."""
        try:
            if isinstance(file_source, bytes):
                docx_stream = io.BytesIO(file_source)
                doc = docx.Document(docx_stream)
            else:
                doc = docx.Document(file_source)

            lines: list[str] = []
            for p in doc.paragraphs:
                if p.text.strip():
                    lines.append(p.text.strip())

            # Extract text embedded in DOCX tables
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_cells:
                        lines.append(" | ".join(row_cells))

            return "\n".join(lines)

        except Exception as e:
            raise ParsingError(f"DOCX parsing failure: {str(e)}") from e
