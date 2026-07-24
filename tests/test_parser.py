"""
Parser Module Unit & Integration Tests
"""

import unittest
import os
from src.parser import ResumeParser, UnsupportedFileTypeError, ParsingError


class TestResumeParser(unittest.TestCase):
    """Test suite for ResumeParser class."""

    def setUp(self):
        self.parser = ResumeParser()

    def test_unsupported_extension_raises_error(self):
        """Verify that unsupported extensions raise UnsupportedFileTypeError."""
        with self.assertRaises(UnsupportedFileTypeError):
            self.parser.parse_file(b"test data", "resume.txt")

    def test_docx_bytes_parsing(self):
        """Verify parsing a valid DOCX byte stream."""
        import docx
        import io

        doc = docx.Document()
        doc.add_paragraph("John Doe - Senior Software Engineer")
        doc.add_paragraph("Skills: Python, Streamlit, Gemini API")

        stream = io.BytesIO()
        doc.save(stream)
        raw_bytes = stream.getvalue()

        result = self.parser.parse_file(raw_bytes, "test_resume.docx")
        self.assertIn("John Doe", result)
        self.assertIn("Python, Streamlit", result)


if __name__ == "__main__":
    unittest.main()
